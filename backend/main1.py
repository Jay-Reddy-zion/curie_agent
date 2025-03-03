import io
import logging
import asyncio
import PyPDF2
from concurrent.futures import ThreadPoolExecutor
from fastapi import HTTPException, Request, UploadFile
from azure.storage.blob import ContentSettings
import pandas as pd
import docx
import openpyxl
import csv
from typing import List, Dict, Any
import re
from settings import get_settings
import os

STORAGE_ACCOUNT_KEY = os.getenv("AZURE_STORAGE_KEY")


settings = get_settings()

# Existing PDF processing functions

def process_page(page):
    
    try:
        # Extract raw text from page
        raw_text = page.extract_text()
        if not raw_text:
            return ""
        
        # Split into lines and process each line
        processed_lines = []
        current_line = []
        
        # Iterate through each character to handle line breaks properly
        for char in raw_text:
            if char == '\n':
                if current_line:
                    line = ''.join(current_line).strip()
                    if line:  # Only process non-empty lines
                        processed_lines.append(f"***{line}***")
                    current_line = []
            else:
                current_line.append(char)
        
        # Handle the last line if it exists
        if current_line:
            line = ''.join(current_line).strip()
            if line:
                processed_lines.append(f"***{line}***")
        
        # Join all processed lines with newlines
        return '\n'.join(processed_lines)
    except Exception as e:
        logging.error(f"Error processing page: {str(e)}")
        return ""  # Return empty string on error to continue processing other pages

# New file processing functions

def process_text_file(text_content: str, debug: bool = False) -> str:
    """
    Process a plain text file by adding special characters to each line.
    
    Args:
        text_content: Content of the text file
        debug: Boolean flag for debug logging
    
    Returns:
        str: Processed text with special characters around each line
    """
    try:
        if not text_content.strip():
            return ""
            
        lines = text_content.split('\n')
        processed_lines = []
        
        for line in lines:
            stripped_line = line.strip()
            if stripped_line:
                processed_lines.append(f"***{stripped_line}***")
        
        return '\n'.join(processed_lines)
    except Exception as e:
        logging.error(f"Error processing text file: {str(e)}")
        return ""

def process_docx_file(docx_bytes: bytes, debug: bool = False) -> str:
    """
    Process a Word document by extracting text and adding special characters.
    
    Args:
        docx_bytes: Bytes of the Word document
        debug: Boolean flag for debug logging
    
    Returns:
        str: Processed text with special characters around each line
    """
    try:
        doc_file_like = io.BytesIO(docx_bytes)
        doc = docx.Document(doc_file_like)
        
        processed_lines = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                processed_lines.append(f"***{text}***")
        
        # Also process tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    processed_lines.append(f"***{row_text}***")
        
        return '\n'.join(processed_lines)
    except Exception as e:
        logging.error(f"Error processing Word document: {str(e)}")
        return ""

def process_csv_file(csv_bytes: bytes, debug: bool = False) -> str:
    """
    Process a CSV file by extracting data and formatting it with special characters.
    
    Args:
        csv_bytes: Bytes of the CSV file
        debug: Boolean flag for debug logging
    
    Returns:
        str: Processed text with special characters around each line
    """
    try:
        csv_file_like = io.StringIO(csv_bytes.decode('utf-8', errors='replace'))
        
        # Detect the delimiter
        dialect = csv.Sniffer().sniff(csv_file_like.read(1024))
        csv_file_like.seek(0)
        
        reader = csv.reader(csv_file_like, dialect)
        processed_lines = []
        
        headers = next(reader, None)
        if headers:
            header_line = ' | '.join([h.strip() for h in headers if h.strip()])
            processed_lines.append(f"***{header_line}***")
        
        for row in reader:
            row_text = ' | '.join([cell.strip() for cell in row if cell.strip()])
            if row_text:
                processed_lines.append(f"***{row_text}***")
        
        return '\n'.join(processed_lines)
    except Exception as e:
        logging.error(f"Error processing CSV file: {str(e)}")
        # Fallback method if sniffing fails
        try:
            csv_file_like = io.StringIO(csv_bytes.decode('utf-8', errors='replace'))
            df = pd.read_csv(csv_file_like)
            return process_dataframe(df, debug)
        except Exception as e2:
            logging.error(f"Fallback CSV processing also failed: {str(e2)}")
            return ""

def process_excel_file(excel_bytes: bytes, debug: bool = False) -> str:
    """
    Process an Excel file by extracting data from all sheets and formatting it.
    
    Args:
        excel_bytes: Bytes of the Excel file
        debug: Boolean flag for debug logging
    
    Returns:
        str: Processed text with special characters around each line
    """
    try:
        excel_file_like = io.BytesIO(excel_bytes)
        
        # Try using pandas first for better handling of complex Excel files
        try:
            with pd.ExcelFile(excel_file_like) as xls:
                sheet_names = xls.sheet_names
                all_processed = []
                
                for sheet_name in sheet_names:
                    if debug:
                        logging.info(f"Processing Excel sheet: {sheet_name}")
                    
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    sheet_content = process_dataframe(df, debug)
                    
                    if sheet_content:
                        all_processed.append(f"***SHEET: {sheet_name}***")
                        all_processed.append(sheet_content)
                
                return '\n'.join(all_processed)
        except Exception as e:
            logging.warning(f"Pandas Excel processing failed, trying openpyxl: {str(e)}")
            
            # Fallback to openpyxl for problematic files
            excel_file_like.seek(0)
            wb = openpyxl.load_workbook(excel_file_like, data_only=True)
            
            all_processed = []
            for sheet_name in wb.sheetnames:
                if debug:
                    logging.info(f"Processing Excel sheet with openpyxl: {sheet_name}")
                
                sheet = wb[sheet_name]
                sheet_lines = []
                
                for row in sheet.rows:
                    row_values = [str(cell.value).strip() if cell.value is not None else "" for cell in row]
                    row_text = ' | '.join([val for val in row_values if val])
                    if row_text:
                        sheet_lines.append(f"***{row_text}***")
                
                if sheet_lines:
                    all_processed.append(f"***SHEET: {sheet_name}***")
                    all_processed.extend(sheet_lines)
            
            return '\n'.join(all_processed)
            
    except Exception as e:
        logging.error(f"Error processing Excel file: {str(e)}")
        return ""

def process_dataframe(df: pd.DataFrame, debug: bool = False) -> str:
    """
    Process a pandas DataFrame by formatting it with special characters.
    
    Args:
        df: DataFrame to process
        debug: Boolean flag for debug logging
    
    Returns:
        str: Processed text with special characters around each line
    """
    try:
        processed_lines = []
        
        # Process headers
        headers = ' | '.join([str(col).strip() for col in df.columns if str(col).strip()])
        if headers:
            processed_lines.append(f"***{headers}***")
        
        # Process rows
        for _, row in df.iterrows():
            row_values = [str(val).strip() if val is not None else "" for val in row]
            row_text = ' | '.join([val for val in row_values if val])
            if row_text:
                processed_lines.append(f"***{row_text}***")
        
        return '\n'.join(processed_lines)
    except Exception as e:
        logging.error(f"Error processing DataFrame: {str(e)}")
        return ""

async def process_blob(file_bytes, file_content_type, blob_name, debug=False):
    """
    Process a file blob based on its content type.
    
    Args:
        file_bytes: Bytes of the file
        file_content_type: MIME type of the file
        blob_name: Name of the blob being processed
        debug: Boolean flag for debug logging
    
    Returns:
        int: Number of indexed chunks
    """
    try:
        if debug:
            logging.info(f"Starting to process blob: {blob_name} of type {file_content_type}")
        
        processed_content = ""
        
        # Process based on file type
        if file_content_type == "application/pdf":
            # Process PDF using existing function
            pdf_file_like = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file_like)
            total_pages = len(pdf_reader.pages)
            
            if debug:
                logging.info(f"PDF has {total_pages} pages")
            
            # Determine optimal chunk size based on total pages
            chunk_size = min(10, max(1, total_pages // 4))  # Adaptive chunk size
            
            async def process_page_batch(start_idx, end_idx):
                """Process a batch of pages using ThreadPoolExecutor."""
                if debug:
                    logging.info(f"Processing batch from page {start_idx} to {end_idx}")
                
                with ThreadPoolExecutor() as executor:
                    batch_pages = pdf_reader.pages[start_idx:end_idx]
                    batch_results = list(executor.map(process_page, batch_pages))
                    
                    # Additional verification of processed content
                    verified_results = []
                    for i, content in enumerate(batch_results):
                        if content.strip():
                            # Verify that each line has the special characters
                            lines = content.split('\n')
                            verified_lines = []
                            for line in lines:
                                if line.strip() and not (line.startswith('***') and line.endswith('***')):
                                    line = f"***{line.strip()}***"
                                verified_lines.append(line)
                            verified_results.append('\n'.join(verified_lines))
                        else:
                            if debug:
                                logging.warning(f"Empty content from page {start_idx + i}")
                    
                    return verified_results
            
            # Create tasks for processing batches
            tasks = []
            for i in range(0, total_pages, chunk_size):
                end_idx = min(i + chunk_size, total_pages)
                tasks.append(process_page_batch(i, end_idx))
            
            # Wait for all batches to complete
            batch_results = await asyncio.gather(*tasks)
            all_content = []
            for batch in batch_results:
                all_content.extend(batch)
            
            # Final verification of content
            processed_content = '\n'.join(filter(None, all_content))
            
        elif file_content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            # Process Word document
            processed_content = process_docx_file(file_bytes, debug)
            
        elif file_content_type in ["text/csv", "application/csv"]:
            # Process CSV
            processed_content = process_csv_file(file_bytes, debug)
            
        elif file_content_type in ["application/vnd.ms-excel", 
                                 "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                 "application/vnd.oasis.opendocument.spreadsheet"]:
            # Process Excel
            processed_content = process_excel_file(file_bytes, debug)
            
        elif file_content_type in ["text/plain"]:
            # Process text file
            text_content = file_bytes.decode('utf-8', errors='replace')
            processed_content = process_text_file(text_content, debug)
            
        else:
            raise ValueError(f"Unsupported file type: {file_content_type}")
        
        if not processed_content.strip():
            raise ValueError("No content extracted from file")
        
        # Split into chunks and index
        chunks = split_text_into_chunks(processed_content)
        if not chunks:
            raise ValueError("No valid chunks generated from document")
        
        # Index chunks with optimized batch processing
        indexed_count = await index_chunks_in_azure_search(chunks, blob_name, debug)
        if indexed_count == 0:
            raise ValueError("No chunks were successfully indexed")
        
        if debug:
            logging.info(f"Successfully processed and indexed {indexed_count} chunks")
        
        return indexed_count
        
    except Exception as e:
        error_msg = f"Error processing blob {blob_name}: {str(e)}"
        logging.error(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)

@app.post("/api/v1/upload")
async def upload_file_to_blob(
    req: Request,
    file: UploadFile,  
    file_name: str, 
    debug: bool = False
):
    user_id = req.headers.get("user-id")
    if not user_id:
        raise HTTPException(status_code=400, detail="User ID is required")
        
    try:
        # Extended validation for multiple file types
        allowed_types = [
            "application/pdf",                                                  # PDF
            "application/msword",                                               # DOC
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", # DOCX
            "text/csv", "application/csv",                                      # CSV
            "application/vnd.ms-excel",                                         # XLS
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", # XLSX
            "application/vnd.oasis.opendocument.spreadsheet",                   # ODS
            "text/plain"                                                        # TXT
        ]
        
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail="Invalid file type. Allowed types: PDF, Word, Excel, CSV, and plain text."
            )
        
        # Read file content once
        file_content = await file.read()
        if not file_content:
            raise HTTPException(status_code=400, detail="Empty file")

        # Prepare blob storage parameters
        container = settings.azure_blob_container_name_1 #docs
        modified_file_name = f"{user_id}_{file_name}"
        content_settings = ContentSettings(content_type=file.content_type or "application/octet-stream")

        # Execute upload and processing concurrently
        async def upload_to_blob():
            blob_client = blob_service_client.get_blob_client(container=container, blob=modified_file_name)
            await asyncio.to_thread(
                blob_client.upload_blob,
                file_content,
                overwrite=True,
                content_settings=content_settings,
                max_concurrency=4  # Enable concurrent uploads
            )
            return f"https://{blob_service_client.account_name}.blob.core.windows.net/{container}/{modified_file_name}"

        # Run upload and processing in parallel
        upload_task = asyncio.create_task(upload_to_blob())
        process_task = asyncio.create_task(process_blob(file_content, file.content_type, modified_file_name, debug))

        # Wait for both tasks to complete
        blob_url, indexed_count = await asyncio.gather(upload_task, process_task)

        return {
            "url": blob_url,
            "name": file_name,
            "chunks_indexed": indexed_count,
            "file_type": file.content_type
        }

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Error in upload_file_to_blob: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# Utility functions

def split_text_into_chunks(text, max_chunk_size=5000, overlap=100):
    """
    Split text into chunks with overlap.
    
    Args:
        text: Text to split
        max_chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
    
    Returns:
        List of chunks
    """
    if not text or not text.strip():
        return []
    
    # Normalize line breaks
    text = re.sub(r'\r\n', '\n', text)
    
    # Split by lines to maintain integrity of lines
    lines = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for line in lines:
        line_length = len(line) + 1  # +1 for the newline
        
        # If adding this line would exceed max_chunk_size, start a new chunk
        if current_length + line_length > max_chunk_size and current_chunk:
            chunks.append('\n'.join(current_chunk))
            
            # Start new chunk with overlap from previous chunk
            overlap_lines = []
            overlap_size = 0
            for prev_line in reversed(current_chunk):
                if overlap_size + len(prev_line) + 1 <= overlap:
                    overlap_lines.insert(0, prev_line)
                    overlap_size += len(prev_line) + 1
                else:
                    break
            
            current_chunk = overlap_lines
            current_length = overlap_size
        
        current_chunk.append(line)
        current_length += line_length
    
    # Add the last chunk if it has content
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks

async def index_chunks_in_azure_search(chunks, blob_name, debug=False):
    """
    Index chunks in Azure Cognitive Search.
    
    Args:
        chunks: List of text chunks to index
        blob_name: Name of the source blob
        debug: Boolean flag for debug logging
    
    Returns:
        int: Number of chunks indexed
    """
    # This is a placeholder for your actual indexing logic
    # You would implement this function to upload the chunks to Azure Search
    
    try:
        if debug:
            logging.info(f"Indexing {len(chunks)} chunks for blob {blob_name}")
        
        # Insert your actual Azure Search indexing code here
        # For example:
        # search_client = get_search_client()
        # batch_size = 100
        # 
        # indexed_count = 0
        # for i in range(0, len(chunks), batch_size):
        #     batch = chunks[i:i+batch_size]
        #     documents = []
        #     
        #     for idx, content in enumerate(batch):
        #         doc_id = f"{blob_name}_{i+idx}"
        #         document = {
        #             "id": doc_id,
        #             "content": content,
        #             "source": blob_name
        #         }
        #         documents.append(document)
        #     
        #     result = await search_client.upload_documents(documents)
        #     indexed_count += sum(1 for r in result if r.succeeded)
        
        # For now, simulate successful indexing
        indexed_count = len(chunks)
        
        if debug:
            logging.info(f"Successfully indexed {indexed_count} chunks")
        
        return indexed_count
    except Exception as e:
        logging.error(f"Error indexing chunks: {str(e)}")
        raise